from pathlib import Path
import os
from PIL import Image
from docx import Document
from docx.shared import Inches
import shutil
import re

Image.MAX_IMAGE_PIXELS = 1000000000 

def get_DPI(filepath):
    img = Image.open(filepath)
    print(img.size)

def goodDPI(img):  
## This will return true for images that are at least 4 x 6. 
## DPI is calculated by pixels/inch, so we find length in inches by pixels/DPI
    bothOver4 = False
    oneOver6 = False

    #If the parameters ever change, change them here
    min_dpi = 300
    min_length = 4
    min_height = 6
    
    pixels = img.size 
    (horizontal, vertical) = pixels
    
    #This allows us the flexibility to have a 4x6 or 6x4 photo
    if horizontal >= min_dpi * min_length and vertical >= min_dpi * min_length:
        bothOver4 = True
    if horizontal >= min_dpi * min_height or vertical >= min_dpi * min_height:
        oneOver6 = True
    
    if bothOver4 == True and oneOver6 == True:
        return True
    else:
        return False

def wordCount(wordFile):
    calledImages = []
    localwordcount = 0
    
    doc = Document(wordFile)
    for par in doc.paragraphs:
            localwordcount = localwordcount + len(re.findall(r'\w+', par.text.lower()))
            callout_num = re.findall(r'insert', par.text.lower()) ##insert\s\w\s\d+
            if len(callout_num) > 0:
                callouts = re.findall(r'\d+\.\d+', par.text)
                for a in callouts:
                    calledImages.append(a)
    
    if len(calledImages) > 0: 
        print(calledImages)
    return localwordcount

def imageReview(filepath):
## This will count images and make sure they meet the standard
    document.add_heading('Image review', 0)
    document.add_paragraph('These images do not meet the required 4x6in @ 300 dpi (1200 x 1800)')
    
    imagecount = 0

    subfolders = [x[0] for x in os.walk(filepath)]                                                                            
    for folder in subfolders:                                                                                            
        files = next(os.walk(folder))[2]                                                                            
        if (len(files) > 0):  
            for file in files:
                if "tif" in file.lower() or "jpg" in file.lower() or "jpeg" in file.lower() or "png" in file.lower():
                    imagecount += 1
                    img = Image.open(folder + "//" + file)
                    if goodDPI(img) == False:   
                        document.add_paragraph(file + ": " + str(img.size) , style='List Bullet')
    document.add_paragraph("Total image count: " + str(imagecount))

def textReview(filepath):
    
    document.add_heading('Text review', 0)
    
    totalWordCount = 0
    bmWordCount = 0
    
    subfolders = [x[0] for x in os.walk(filepath)]                                                                            
    for folder in subfolders:                                                                                            
        files = next(os.walk(folder))[2]                                                                            
        if (len(files) > 0):  
            for file in files:
                if ".doc" in file.lower():
                    if "bm" in file.lower():
                        bmWordCount += wordCount(folder + "//" + file)
                    elif "fm" in file.lower():
                        document.add_paragraph('FM word count: ' + str(wordCount(folder + "//" + file)))
                    else:
                        totalWordCount += wordCount(folder + "//" + file)
                        
    document.add_paragraph('BM word count: ' + str(bmWordCount))
    document.add_paragraph('Body word count: ' + str(totalWordCount))            

def transmittal():                    
    document = Document()
    path = input("Filepath: ")

    imageReview(path)
    textReview(path)

    document.save('TM Notes.docx')
    shutil.move('TM Notes.docx', path)


#get_DPI("/Users/karenli/Box/Books (rutgerspress2)/SpringSummer_21/Freeland, David/Interior/TM Art/1.1.jpg") 
#How to tell if a photo is missing?
document = Document()
path = input("Filepath: ")

imageReview(path)
textReview(path)

document.save('TM Notes.docx')
shutil.move('TM Notes.docx', path)

##If you have an error like the following "Package not found at '%s'" % pkg_file docx.opc.exceptions.PackageNotFoundError: Package not found at 'filepath'
## Copy the files and paste them in a new folder
## The program might be caught on hidden unsaved versions of files. 
## Also try closing word. Again, ghost files